"""
Document extraction functions using Google Document AI for OCR.
"""
import os
import logging
from pathlib import Path
from google.cloud import documentai

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Google Document AI configuration
PROJECT_ID = "704440646290"
LOCATION = "us"
PROCESSOR_ID = "8ecc3543a9209b03"
CREDENTIALS_PATH = r"C:\Users\ChristopherCato\OneDrive - clarity-dx.com\Documents\Bill_Review_INTERNAL\OCR\bill-review-ocr-529ae52caf66.json"

# Initialize Document AI Client
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH
documentai_client = None

def initialize_documentai():
    """Initialize the Document AI client."""
    global documentai_client
    try:
        documentai_client = documentai.DocumentProcessorServiceClient()
        logger.info("Document AI client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Document AI client: {str(e)}")
        raise

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using Google Document AI.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as string
    """
    global documentai_client
    if documentai_client is None:
        initialize_documentai()
        
    try:
        logger.info(f"Processing PDF with Document AI: {file_path}")
        
        # Read the PDF file
        with open(file_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
            
        # Create a raw document request
        raw_document = documentai.RawDocument(content=pdf_data, mime_type="application/pdf")
        
        # Build processor resource name
        processor_name = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"
        
        # Create OCR request
        request = documentai.ProcessRequest(name=processor_name, raw_document=raw_document)
        
        # Process document
        result = documentai_client.process_document(request=request)
        
        # Extract text from result
        return result.document.text
    except Exception as e:
        logger.error(f"Error extracting PDF with Document AI {file_path}: {str(e)}")
        return f"[Error extracting PDF: {str(e)}]"

def extract_text_from_docx(file_path):
    """Extract text from Word documents."""
    try:
        import docx
        
        text = ""
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
        return f"[Error extracting DOCX: {str(e)}]"

def extract_text_from_email(file_path):
    """Extract text and metadata from email files."""
    try:
        import email
        from email import policy
        from email.parser import BytesParser
        
        with open(file_path, 'rb') as fp:
            msg = BytesParser(policy=policy.default).parse(fp)
        
        # Extract basic metadata
        metadata = {
            "from": str(msg.get('from', '')),
            "to": str(msg.get('to', '')),
            "subject": str(msg.get('subject', '')),
            "date": str(msg.get('date', ''))
        }
        
        # Get body content
        body = ""
        if msg.is_multipart():
            for part in msg.iter_parts():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    body += part.get_content() + "\n"
        else:
            if msg.get_content_type() == "text/plain":
                body = msg.get_content()
        
        return f"""
Email Metadata:
From: {metadata['from']}
To: {metadata['to']}
Subject: {metadata['subject']}
Date: {metadata['date']}

Email Body:
{body}
"""
    except Exception as e:
        logger.error(f"Error extracting email {file_path}: {str(e)}")
        return f"[Error extracting email: {str(e)}]"

def extract_text_from_txt(file_path):
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text file {file_path}: {str(e)}")
        return f"[Error extracting text file: {str(e)}]"

def extract_text_from_image(file_path):
    """Extract text from images using Google Document AI."""
    try:
        logger.info(f"Processing image with Document AI: {file_path}")
        
        # Determine the mime type based on file extension
        extension = Path(file_path).suffix.lower()
        mime_type = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".tiff": "image/tiff",
            ".bmp": "image/bmp"
        }.get(extension, "image/jpeg")
        
        # Read the image file
        with open(file_path, "rb") as image_file:
            image_data = image_file.read()
            
        # Create a raw document request
        raw_document = documentai.RawDocument(content=image_data, mime_type=mime_type)
        
        # Build processor resource name
        processor_name = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"
        
        # Create OCR request
        request = documentai.ProcessRequest(name=processor_name, raw_document=raw_document)
        
        # Process document
        result = documentai_client.process_document(request=request)
        
        # Extract text from result
        return result.document.text
    except Exception as e:
        logger.error(f"Error extracting image with Document AI {file_path}: {str(e)}")
        return f"[Error extracting image: {str(e)}]"

def extract_text(file_path):
    """
    Extract text from a file based on its extension using Document AI for PDF/images.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text as string
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    # Call appropriate extractor based on file extension
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        return extract_text_from_image(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension == '.eml':
        return extract_text_from_email(file_path)
    elif extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        message = f"Unsupported file type: {extension}"
        logger.warning(message)
        return f"[{message}]"